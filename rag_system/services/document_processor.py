# rag_system/services/document_processor.py
import os
from typing import List, Dict, Any
from django.conf import settings
from ..models import Document, DocumentChunk


class DocumentProcessor:
    """
    Reads PDF, DOCX, TXT, MD files and plain text strings into the RAG system.
    PDF reading uses pypdf (in requirements.txt).
    DOCX reading uses python-docx (in requirements.txt).
    """

    def __init__(self):
        self.chunk_size    = settings.RAG_CONFIG['CHUNK_SIZE']
        self.chunk_overlap = settings.RAG_CONFIG['CHUNK_OVERLAP']

    # ── Public file-reading entry point ──────────────────────────────────────

    def read_file(self, file_path: str) -> str:
        """
        Extract text from a file on disk.
        Supported: .pdf  .docx  .doc  .txt  .md
        Raises ValueError with a clear message for unsupported types.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        readers = {
            '.pdf':    self._read_pdf,
            '.docx':   self._read_docx,
            '.doc':    self._read_docx,
            '.txt':    self._read_text,
            '.md':     self._read_text,
            '.py':     self._read_text,
            '.ipynb':  self._read_notebook,
        }

        if ext not in readers:
            raise ValueError(
                f"Unsupported file type: '{ext}'. "
                f"Supported: {', '.join(readers.keys())}"
            )

        text = readers[ext](file_path)
        if not text or not text.strip():
            raise ValueError(
                f"No text could be extracted from {os.path.basename(file_path)}. "
                "If it's a scanned PDF, the text is in images and cannot be read. "
                "Try copying the text manually into the 'description' or 'inline notes' field."
            )
        return text

    # ── Format readers ────────────────────────────────────────────────────────

    def _read_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pypdf (modern, maintained)."""
        # FIX 4: use pypdf not the deprecated PyPDF2
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError(
                "pypdf is required for PDF reading. "
                "Run: pip install pypdf"
            )
        reader = PdfReader(file_path)
        pages  = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ''
                if text.strip():
                    pages.append(text)
            except Exception as e:
                print(f"  Warning: could not extract page {i + 1}: {e}")

        if not pages:
            raise ValueError(
                "No text could be extracted from this PDF. "
                "It may be scanned or image-based. "
                "Try copying the text manually into the description field."
            )
        return '\n\n'.join(pages)

    def _read_docx(self, file_path: str) -> str:
        """Extract text from Word documents (paragraphs + table cells)."""
        try:
            import docx as python_docx
        except ImportError:
            raise ImportError(
                "python-docx is required for Word files. "
                "Run: pip install python-docx"
            )
        doc   = python_docx.Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = '  |  '.join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return '\n'.join(parts)

    def _read_text(self, file_path: str) -> str:
        """Read plain text, Markdown, or Python source code."""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()

    def _read_notebook(self, file_path: str) -> str:
        """Extract code and markdown cells from a Jupyter notebook (.ipynb)."""
        import json
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            nb = json.load(f)
        parts = []
        for cell in nb.get('cells', []):
            cell_type = cell.get('cell_type', '')
            source    = ''.join(cell.get('source', []))
            if not source.strip():
                continue
            if cell_type == 'markdown':
                parts.append(source)
            elif cell_type == 'code':
                parts.append(f"```python\n{source}\n```")
        return '\n\n'.join(parts)

    # ── Chunking ──────────────────────────────────────────────────────────────

    def chunk_text(self, text: str, title: str = '') -> list:
        """Split text into overlapping word-based chunks."""
        words  = text.split()
        step   = max(1, self.chunk_size - self.chunk_overlap)
        chunks = []
        for i, start in enumerate(range(0, len(words), step)):
            text_chunk = ' '.join(words[start: start + self.chunk_size])
            if not text_chunk.strip():
                continue
            chunks.append({
                'content': text_chunk,
                'metadata': {
                    'chunk_index':  i,
                    'document_title': title,
                    'start_word':   start,
                    'end_word':     min(start + self.chunk_size, len(words)),
                }
            })
        return chunks

    # ── High-level helpers ────────────────────────────────────────────────────

    def process_document(self, file_path: str, document_type: str,
                         title: str = None) -> Document:
        """
        Read a file, chunk it, embed it, save everything to the database.
        Returns the created Document instance.
        """
        if title is None:
            title = os.path.splitext(os.path.basename(file_path))[0].replace('_', ' ').title()

        content = self.read_file(file_path)

        # Delete any existing document from the same source so we don't duplicate
        Document.objects.filter(source=file_path).delete()

        doc = Document.objects.create(
            title=title,
            content=content,
            document_type=document_type,
            source=file_path,
            metadata={'file_size': len(content), 'file_path': file_path},
        )
        chunks = self.chunk_text(content, title)
        chunk_objs = [
            DocumentChunk(
                document=doc,
                content=c['content'],
                chunk_index=c['metadata']['chunk_index'],
                metadata={**c['metadata'], 'document_type': document_type},
            )
            for c in chunks
        ]
        DocumentChunk.objects.bulk_create(chunk_objs)
        return doc

    def ingest_text(self, content: str, title: str,
                    document_type: str, source: str) -> Document:
        """
        Ingest a plain-text string. Creates/replaces the Document and
        all its chunks. Used by ingest_from_db for inline text fields.
        """
        Document.objects.filter(source=source).delete()
        doc = Document.objects.create(
            title=title,
            content=content,
            document_type=document_type,
            source=source,
            metadata={'source': source},
        )
        chunk_objs = [
            DocumentChunk(
                document=doc,
                content=c['content'],
                chunk_index=c['metadata']['chunk_index'],
                metadata={**c['metadata'], 'document_type': document_type},
            )
            for c in self.chunk_text(content, title)
        ]
        DocumentChunk.objects.bulk_create(chunk_objs)
        return doc
